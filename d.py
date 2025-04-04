import streamlit as st
import pandas as pd
import nltk
import re
import io
import base64
from newspaper import Article
import requests
from streamlit_lottie import st_lottie
from nltk.corpus import stopwords
from nltk.sentiment.vader import SentimentIntensityAnalyzer
from nltk.tokenize import word_tokenize, sent_tokenize
from fpdf import FPDF
import docx

# Download required NLTK data
nltk.download('punkt')
nltk.download('stopwords')
nltk.download('vader_lexicon')

# Function to load Lottie animations
def load_lottieurl(url: str):
    r = requests.get(url)
    if r.status_code != 200:
        return None
    return r.json()

# Function to extract article text from URL
def get_article_text(url):
    article = Article(url)
    article.download()
    article.parse()
    return article.title, article.text

# Function to clean text
def clean_text(text, stopword_file=None):
    stop_words = set(stopwords.words('english'))
    if stopword_file:
        try:
            with open(stopword_file, 'r') as f:
                for line in f:
                    stop_words.add(line.strip())
        except FileNotFoundError:
            st.warning(f"Stopwords file {stopword_file} not found. Using default stopwords.")
    
    word_tokens = word_tokenize(text)
    filtered_text = [w for w in word_tokens if not w in stop_words]
    return filtered_text

# Function to check if word is complex
def is_complex(word):
    if len(word) > 7:
        return True
    else:
        return False

# Function to calculate readability metrics
def get_readability(text):
    sentences = sent_tokenize(text)
    words = word_tokenize(text)
    
    if len(sentences) == 0:
        return 0, 0, 0, 0
    
    avg_sentence_length = len(words) / len(sentences)
    
    if len(words) == 0:
        return avg_sentence_length, 0, 0, 0
    
    complex_words = [w for w in words if is_complex(w)]
    percent_complex_words = len(complex_words) / len(words) * 100
    fog_index = 0.4 * (avg_sentence_length + percent_complex_words)
    return avg_sentence_length, percent_complex_words, fog_index, len(complex_words)

# Function to count words
def get_word_count(text):
    words = word_tokenize(text)
    words = [w for w in words if w.isalpha()]
    return len(words)

# Function to count syllables in a word
def count_syllables(word):
    vowels = "aeiouy"
    word = word.lower()
    if word[0] in vowels:
        count = 1
    else:
        count = 0
    for i in range(1, len(word)):
        if word[i] in vowels and word[i - 1] not in vowels:
            count += 1
    if word.endswith("e"):
        count -= 1
    if word.endswith("le") and len(word) > 2 and word[-3] not in vowels:
        count += 1
    if count == 0:
        count += 1
    return count

# Function to get syllable count for all words
def get_syllable_count(text):
    words = word_tokenize(text)
    words = [w for w in words if w.isalpha()]
    total_syllables = sum([count_syllables(w) for w in words])
    avg_syllables = total_syllables / len(words) if len(words) > 0 else 0
    return avg_syllables

# Function to count personal pronouns
def get_personal_pronouns(text):
    pattern = r'\b[Ii]|[Ww]e|[Mm]y|[Oo]urs|[Uu]s\b'
    personal_pronouns = re.findall(pattern, text)
    personal_pronouns = [w for w in personal_pronouns if w.lower() != 'us']
    return len(personal_pronouns)

# Function to calculate average word length
def get_avg_word_length(text):
    words = word_tokenize(text)
    words = [w for w in words if w.isalpha()]
    if len(words) == 0:
        return 0
    total_chars = sum([len(w) for w in words])
    avg_word_length = total_chars / len(words)
    return avg_word_length

# Function to get sentiment scores using both custom dictionaries and VADER
def get_sentiment(text, pos_dict_file=None, neg_dict_file=None, stopword_file=None):
    # VADER sentiment analysis
    sid = SentimentIntensityAnalyzer()
    vader_scores = sid.polarity_scores(text)
    
    # Custom dictionary sentiment analysis (if files are provided)
    pos_score, neg_score, polarity_score, subjectivity_score = 0, 0, 0, 0
    
    if pos_dict_file and neg_dict_file and stopword_file:
        try:
            cleaned_text = clean_text(text, stopword_file)
            pos_dict = set()
            neg_dict = set()
            
            try:
                with open(pos_dict_file, 'r') as f:
                    for line in f:
                        pos_dict.add(line.strip())
            except FileNotFoundError:
                st.warning(f"Positive dictionary file {pos_dict_file} not found.")
                
            try:
                with open(neg_dict_file, 'r') as f:
                    for line in f:
                        neg_dict.add(line.strip())
            except FileNotFoundError:
                st.warning(f"Negative dictionary file {neg_dict_file} not found.")
                
            pos_score = sum([1 for word in cleaned_text if word in pos_dict])
            neg_score = sum([1 for word in cleaned_text if word in neg_dict])
            
            if pos_score + neg_score > 0:
                polarity_score = (pos_score - neg_score) / (pos_score + neg_score)
                subjectivity_score = (pos_score + neg_score) / len(cleaned_text) if len(cleaned_text) > 0 else 0
            else:
                polarity_score = 0
                subjectivity_score = 0
        except Exception as e:
            st.error(f"Error in custom sentiment analysis: {e}")
    
    return pos_score, neg_score, polarity_score, subjectivity_score, vader_scores

# Set page title
st.set_page_config(page_title="Article Sentiment Analysis", layout="wide")

# Display Lottie animation
lottie_url_hello = "https://lottie.host/0d49d388-9acb-492c-92c2-8dad820db057/R3WmiFyHXU.json"
lottie_hello = load_lottieurl(lottie_url_hello)
st_lottie(lottie_hello, height=200)

# Page title
st.title('Article Sentiment Analysis')

# Input URL
url = st.text_input('Enter URL of the article:')

if url:
    with st.spinner('Extracting article content...'):
        try:
            # Extract article
            title, text = get_article_text(url)
            
            # Save article to file - FIXED SYNTAX ERROR HERE
            pattern = r'[^\w\s]'
            sanitized_title = re.sub(pattern, '', title)
            file_name = f"{sanitized_title[:50]}.txt"
            with open(file_name, 'w', encoding='utf-8') as f:
                f.write(text)
                
            st.success('Article extracted successfully!')
            
            # Display article information
            with st.expander("View Article Content"):
                st.subheader(title)
                st.write(text)
            
            # Perform analysis
            with st.spinner('Analyzing content...'):
                # Get sentiment scores
                pos_score, neg_score, polarity_score, subjectivity_score, vader_scores = get_sentiment(
                    text, 'positive.txt', 'negative.txt', 'stopwords.txt'
                )
                
                # Get readability metrics
                avg_sentence_length, percent_complex_words, fog_index, complex_word_count = get_readability(text)
                
                # Get other metrics
                word_count = get_word_count(text)
                avg_syllables = get_syllable_count(text)
                personal_pronouns_count = get_personal_pronouns(text)
                avg_word_length = get_avg_word_length(text)
                
                # Create results dictionary
                result = {
                    'URL': url,
                    'Article Title': title,
                    'Word Count': word_count,
                    'VADER Positive Score': vader_scores['pos'],
                    'VADER Negative Score': vader_scores['neg'],
                    'VADER Neutral Score': vader_scores['neu'],
                    'VADER Compound Score': vader_scores['compound'],
                    'Custom Positive Score': pos_score,
                    'Custom Negative Score': neg_score,
                    'Custom Polarity Score': polarity_score,
                    'Subjectivity Score': subjectivity_score,
                    'Average Sentence Length': avg_sentence_length,
                    'Percentage of Complex Words': percent_complex_words,
                    'Fog Index': fog_index,
                    'Complex Word Count': complex_word_count,
                    'Average Syllables Per Word': avg_syllables,
                    'Personal Pronouns Count': personal_pronouns_count,
                    'Average Word Length': avg_word_length
                }
                
                # Create DataFrame and display results
                result_df = pd.DataFrame([result])
                
                # Visualization section
                st.subheader('Analysis Results')
                
                # Create two columns for the dashboard
                col1, col2 = st.columns(2)
                
                with col1:
                    st.metric("Word Count", word_count)
                    st.metric("Average Sentence Length", f"{avg_sentence_length:.2f}")
                    st.metric("Fog Index (Readability)", f"{fog_index:.2f}")
                
                with col2:
                    sentiment_label = "Positive" if vader_scores['compound'] > 0.05 else "Negative" if vader_scores['compound'] < -0.05 else "Neutral"
                    sentiment_color = "green" if sentiment_label == "Positive" else "red" if sentiment_label == "Negative" else "blue"
                    st.metric("Overall Sentiment", sentiment_label)
                    st.metric("VADER Compound Score", f"{vader_scores['compound']:.2f}")
                
                # Display detailed results
                with st.expander("View Detailed Analysis Results"):
                    st.dataframe(result_df)
                
                # Save results to Excel
                excel_buffer = io.BytesIO()
                with pd.ExcelWriter(excel_buffer, engine='xlsxwriter') as writer:
                    result_df.to_excel(writer, sheet_name='Analysis Results', index=False)
                excel_data = excel_buffer.getvalue()
                
                # Download options
                st.subheader("Download Options")
                
                col1, col2 = st.columns(2)
                
                with col1:
                    # Download Excel results - FIXED SYNTAX ERROR HERE
                    excel_filename = f"{sanitized_title[:30]}_analysis.xlsx"
                    st.download_button(
                        label="Download Analysis Results (Excel)",
                        data=excel_data,
                        file_name=excel_filename,
                        mime="application/vnd.ms-excel"
                    )
                
                with col2:
                    # Select format for article download
                    download_format = st.selectbox("Select Article Download Format", ["PDF", "Text", "DOCX"])
                    
                    if download_format == "PDF":
                        pdf = FPDF()
                        pdf.add_page()
                        pdf.set_font("Arial", size=12)
                        pdf.cell(200, 10, txt=title, ln=True, align='C')
                        pdf.ln(10)
                        
                        # Split text into smaller chunks to avoid issues with long text
                        chunks = [text[i:i+1000] for i in range(0, len(text), 1000)]
                        for chunk in chunks:
                            pdf.multi_cell(0, 10, txt=chunk)
                            
                        pdf_output = io.BytesIO()
                        pdf.output(pdf_output)
                        pdf_data = pdf_output.getvalue()
                        
                        # FIXED SYNTAX ERROR HERE
                        pdf_filename = f"{sanitized_title[:30]}.pdf"
                        st.download_button(
                            label=f"Download Article as {download_format}",
                            data=pdf_data,
                            file_name=pdf_filename,
                            mime="application/pdf"
                        )
                    
                    elif download_format == "Text":
                        # FIXED SYNTAX ERROR HERE
                        txt_filename = f"{sanitized_title[:30]}.txt"
                        st.download_button(
                            label=f"Download Article as {download_format}",
                            data=text,
                            file_name=txt_filename,
                            mime="text/plain"
                        )
                    
                    elif download_format == "DOCX":
                        doc = docx.Document()
                        doc.add_heading(title, 0)
                        doc.add_paragraph(text)
                        
                        docx_output = io.BytesIO()
                        doc.save(docx_output)
                        docx_data = docx_output.getvalue()
                        
                        # FIXED SYNTAX ERROR HERE
                        docx_filename = f"{sanitized_title[:30]}.docx"
                        st.download_button(
                            label=f"Download Article as {download_format}",
                            data=docx_data,
                            file_name=docx_filename,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
        
        except Exception as e:
            st.error(f"Error processing URL: {str(e)}")
            st.info("Please check if the URL is valid and the article is accessible.")

# Add information about the app
with st.expander("About this app"):
    st.markdown("""
    # Article Sentiment Analysis App
    
    This application extracts text from articles at a given URL and performs sentiment analysis and readability assessment.
    
    ## Features:
    - Article extraction using the newspaper library
    - Sentiment analysis using both VADER and custom dictionaries
    - Readability metrics calculation (Fog Index)
    - Text statistics (word count, sentence length, etc.)
    - Download options in multiple formats
    
    ## How to use:
    1. Enter a valid URL in the input field
    2. Wait for the analysis to complete
    3. View the results and download in your preferred format
    """)
